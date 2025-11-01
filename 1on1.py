import streamlit as st
import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime
import json
import os
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
import re

# 메인 컨텐츠 최대 너비 제한 (우측 영역)
st.markdown(
    """
    <style>
    .main .block-container {
        max-width: 800px;
        margin-left: auto;
        margin-right: auto;
        padding-left: 1rem;
        padding-right: 1rem;
    }
    
    /* 성과 코칭 피드백, 성장 코칭 피드백 타이틀 폰트 크기 확대 */
    /* Streamlit expander 헤더 스타일 - 기본 */
    div[data-testid="stExpander"] details summary,
    div[data-testid="stExpander"] summary,
    .streamlit-expanderHeader,
    div[data-testid="stExpander"] details summary p,
    div[data-testid="stExpander"] details summary label,
    div[data-testid="stExpander"] details summary div,
    div[data-testid="stExpander"] details summary span {
        font-size: 2.2rem !important;
        font-weight: 700 !important;
        line-height: 1.6 !important;
    }
    
    /* expander 내부 텍스트 요소들도 크게 */
    div[data-testid="stExpander"] summary * {
        font-size: 2.2rem !important;
        font-weight: 700 !important;
    }
    
    /* 성과 코칭 피드백 - 파랑 계열 색상 */
    div[data-testid="stExpander"]:has(summary:contains("성과")) details summary,
    div[data-testid="stExpander"]:has(summary:contains("성과")) summary,
    div[data-testid="stExpander"]:has(summary:contains("성과")) summary * {
        color: #1E88E5 !important; /* 파랑색 - 성과 */
    }
    
    /* 성장 코칭 피드백 - 초록 계열 색상 */
    div[data-testid="stExpander"]:has(summary:contains("성장")) details summary,
    div[data-testid="stExpander"]:has(summary:contains("성장")) summary,
    div[data-testid="stExpander"]:has(summary:contains("성장")) summary * {
        color: #27AE60 !important; /* 초록색 - 성장 */
    }
    </style>
    """,
    unsafe_allow_html=True
)

# JavaScript로 색상 적용 (CSS :has 선택자 미지원 브라우저 대비 및 동적 콘텐츠 대응)
st.markdown("""
<script>
function applyExpanderColors() {
    const expanders = document.querySelectorAll('div[data-testid="stExpander"] summary');
    expanders.forEach(function(summary) {
        const text = summary.textContent || summary.innerText || '';
        if (text.includes('성과')) {
            summary.style.color = '#1E88E5';
            // summary 내부의 모든 요소에도 색상 적용
            const elements = summary.querySelectorAll('*');
            elements.forEach(function(el) {
                el.style.color = '#1E88E5';
            });
        } else if (text.includes('성장')) {
            summary.style.color = '#27AE60';
            // summary 내부의 모든 요소에도 색상 적용
            const elements = summary.querySelectorAll('*');
            elements.forEach(function(el) {
                el.style.color = '#27AE60';
            });
        }
    });
}

// DOM 로드 후 실행
if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', applyExpanderColors);
} else {
    applyExpanderColors();
}

// Streamlit이 콘텐츠를 업데이트할 때마다 실행
const observer = new MutationObserver(applyExpanderColors);
observer.observe(document.body, { childList: true, subtree: true });

// 주기적으로 체크 (Streamlit의 동적 콘텐츠 업데이트 대응)
setInterval(applyExpanderColors, 500);
</script>
""", unsafe_allow_html=True)

# Google Sheets 연동을 위한 설정
SCOPE = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

# Gemini API 키 설정 (Streamlit secrets 사용)
GEMINI_API_KEY = None
try:
    if hasattr(st, "secrets"):
        gemini_sec = st.secrets.get("gemini", {})
        GEMINI_API_KEY = gemini_sec.get("api_key")
except Exception:
    pass

if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
    except Exception:
        pass

# 1on1 코칭 스프레드시트 ID는 함수에서 동적으로 가져옵니다 (안전성을 위해)
def get_oneon1_spreadsheet_id():
    """1on1 코칭 스프레드시트 ID를 안전하게 가져옵니다."""
    try:
        if hasattr(st, "secrets"):
            secrets_obj = getattr(st, "secrets", None)
            if secrets_obj and hasattr(secrets_obj, "get"):
                try:
                    google_sec = secrets_obj.get("google", {})
                    if google_sec and isinstance(google_sec, dict):
                        spreadsheet_id = google_sec.get("oneon1_spreadsheet_id")
                        if spreadsheet_id:
                            return str(spreadsheet_id)
                except (AttributeError, TypeError):
                    pass
    except Exception:
        pass
    
    # Streamlit secrets에서 가져오기 (추가 시크릿 위치 확인)
    try:
        if hasattr(st, "secrets"):
            # 추가 secrets 위치 확인
            direct_id = st.secrets.get("oneon1_spreadsheet_id") or st.secrets.get("ONEOONE_SPREADSHEET_ID")
            if direct_id:
                return str(direct_id)
    except Exception:
        pass
    
    return ""

# 모듈 로드 시점에는 None으로 초기화, 실제 사용 시점에 함수 호출
ONEOONE_SPREADSHEET_ID = None

def get_google_sheets_client():
    """Google Sheets 클라이언트를 반환합니다."""
    try:
        # 방법 0: Streamlit secrets에서 서비스 계정 정보 읽기 (최우선)
        try:
            google_sec = st.secrets.get("google") if hasattr(st, "secrets") else None
        except Exception:
            google_sec = None

        if google_sec:
            # 0-a) 전체 JSON 문자열 저장한 경우: google.credentials_json
            if "credentials_json" in google_sec and google_sec["credentials_json"]:
                raw = google_sec["credentials_json"]
                creds_info = json.loads(raw) if isinstance(raw, str) else dict(raw)
                creds = Credentials.from_service_account_info(creds_info, scopes=SCOPE)
                return gspread.authorize(creds)

            # 0-b) TOML로 키를 중첩(dict) 저장한 경우: google.service_account = { ... }
            if "service_account" in google_sec and google_sec["service_account"]:
                creds_info = dict(google_sec["service_account"])  # MappingProxyType 대응
                creds = Credentials.from_service_account_info(creds_info, scopes=SCOPE)
                return gspread.authorize(creds)

        # 방법 1: Streamlit secrets에서 직접 읽기 (추가 위치 확인)
        try:
            if hasattr(st, "secrets"):
                direct_creds = st.secrets.get("GOOGLE_CREDENTIALS_JSON") or st.secrets.get("google_credentials_json")
                if direct_creds:
                    try:
                        creds_info = json.loads(direct_creds) if isinstance(direct_creds, str) else dict(direct_creds)
                        creds = Credentials.from_service_account_info(creds_info, scopes=SCOPE)
                        return gspread.authorize(creds)
                    except json.JSONDecodeError:
                        st.error("Streamlit secrets의 GOOGLE_CREDENTIALS_JSON이 올바른 JSON 형식이 아닙니다.")
                        return None
        except Exception:
            pass
        
        # 방법 2: 서비스 계정 JSON 파일 읽기
        service_account_file = "service_account.json"
        if os.path.exists(service_account_file):
            creds = Credentials.from_service_account_file(service_account_file, scopes=SCOPE)
            return gspread.authorize(creds)
        
        # 방법 3: 세션 상태에서 서비스 계정 정보 읽기
        if 'google_credentials' in st.session_state and st.session_state.google_credentials:
            try:
                creds_info = json.loads(st.session_state.google_credentials)
                creds = Credentials.from_service_account_info(creds_info, scopes=SCOPE)
                return gspread.authorize(creds)
            except json.JSONDecodeError:
                st.error("저장된 서비스 계정 정보가 올바른 JSON 형식이 아닙니다.")
                return None
        
        # 인증 정보가 없는 경우
        st.warning("Google Sheets 연동을 위해 서비스 계정 인증 정보가 필요합니다.")
        return None
        
    except Exception as e:
        st.error(f"Google Sheets 연동 오류: {e}")
        return None

def get_oneon1_dataframe():
    """1on1 코칭 데이터를 Google Sheets에서 가져옵니다."""
    try:
        spreadsheet_id = get_oneon1_spreadsheet_id()
        if not spreadsheet_id:
            return None
        
        client = get_google_sheets_client()
        if not client:
            return None
        
        spreadsheet = client.open_by_key(spreadsheet_id)
        worksheet = spreadsheet.worksheet("Sheet1")
        records = worksheet.get_all_records()
        return pd.DataFrame(records)
    except Exception as e:
        st.error(f"1on1 데이터 로드 오류: {e}")
        return None

def save_oneon1_record(data):
    """1on1 코칭 기록을 Google Sheets에 저장합니다."""
    try:
        spreadsheet_id = get_oneon1_spreadsheet_id()
        if not spreadsheet_id:
            st.error("1on1 스프레드시트 ID가 설정되지 않았습니다.")
            return False
        
        client = get_google_sheets_client()
        if not client:
            return False
        
        spreadsheet = client.open_by_key(spreadsheet_id)
        worksheet = spreadsheet.worksheet("Sheet1")
        
        # 데이터를 행으로 추가
        worksheet.append_row(data)
        return True
    except Exception as e:
        st.error(f"1on1 기록 저장 오류: {e}")
        return False

def render_oneon1_form():
    """1on1 코칭 기록 양식을 렌더링합니다."""
    st.subheader("📝 1on1 코칭 기록 작성")
    
    with st.form("oneon1_form", clear_on_submit=False):
        # 코칭 날짜
        coaching_date = st.date_input(
            "코칭 날짜",
            value=datetime.now().date(),
            help="1on1 코칭을 진행한 날짜를 선택하세요"
        )
        
        # 참여자 정보
        col1, col2 = st.columns(2)
        with col1:
            coach_name = st.text_input(
                "코치 이름",
                value=st.session_state.user_info.get('name', '') if st.session_state.get('user_info') else '',
                help="코치(리더) 이름을 입력하세요"
            )
        with col2:
            coachee_name = st.text_input(
                "코치이 이름",
                placeholder="코칭 받는 사람의 이름",
                help="코치이(구성원) 이름을 입력하세요"
            )
        
        # 코칭 주제
        coaching_topic = st.text_input(
            "코칭 주제",
            placeholder="예: 업무 진행 상황, 목표 달성도, 성장 포인트 등",
            help="이번 코칭에서 다룬 주제를 입력하세요"
        )
        
        # 주요 내용
        main_content = st.text_area(
            "주요 내용",
            height=150,
            placeholder="코칭에서 다룬 주요 내용을 입력하세요",
            help="코칭 중 나눈 대화의 핵심 내용을 기록하세요"
        )
        
        # 액션 아이템
        action_items = st.text_area(
            "액션 아이템",
            height=100,
            placeholder="예: * 다음주까지 문서 작성 완료\n* 월말까지 프로젝트 계획 수립",
            help="다음 코칭까지 해야 할 일들을 입력하세요"
        )
        
        # 다음 미팅 일정
        next_meeting = st.date_input(
            "다음 미팅 일정",
            value=None,
            help="다음 1on1 코칭 예정일을 선택하세요 (선택사항)"
        )
        
        # 코칭 평가
        st.markdown("### 코칭 평가")
        coaching_quality = st.slider(
            "코칭 질 평가",
            min_value=1,
            max_value=5,
            value=3,
            help="이번 코칭의 질을 1-5점으로 평가하세요"
        )
        
        notes = st.text_area(
            "기타 메모",
            height=100,
            placeholder="추가로 기록할 사항이 있으면 입력하세요",
            help="기타 메모 사항을 입력하세요 (선택사항)"
        )
        
        # 제출 버튼
        submitted = st.form_submit_button(
            "📤 기록 저장하기",
            use_container_width=True,
            type="primary"
        )
        
        if submitted:
            # 필수 필드 검증
            if not coach_name.strip():
                st.error("코치 이름을 입력해주세요.")
                return False
            if not coachee_name.strip():
                st.error("코치이 이름을 입력해주세요.")
                return False
            if not coaching_topic.strip():
                st.error("코칭 주제를 입력해주세요.")
                return False
            
            # 데이터 준비
            now = datetime.now()
            timestamp = now.strftime("%Y. %m. %d %p %I:%M:%S").replace("AM", "오전").replace("PM", "오후")
            
            data = [
                timestamp,
                str(coaching_date),
                coach_name,
                coachee_name,
                coaching_topic,
                main_content,
                action_items,
                str(next_meeting) if next_meeting else "",
                coaching_quality,
                notes
            ]
            
            # 저장
            if save_oneon1_record(data):
                st.success("✅ 1on1 코칭 기록이 성공적으로 저장되었습니다!")
                st.rerun()
            else:
                st.error("❌ 저장 중 오류가 발생했습니다. 다시 시도해주세요.")
            
            return True
    
    return False

def render_oneon1_history():
    """1on1 코칭 기록 내역을 렌더링합니다."""
    st.subheader("📚 1on1 코칭 기록 내역")
    
    # 데이터 가져오기
    df = get_oneon1_dataframe()
    
    if df is None or df.empty:
        st.info("아직 기록된 1on1 코칭 내역이 없습니다.")
        return
    
    # 사용자 필터링 (로그인된 사용자만 자신의 기록 보기)
    user_name = None
    if st.session_state.get('logged_in') and st.session_state.get('user_info'):
        user_name = st.session_state.user_info.get('name', '')
    
    if user_name:
        # 코치 또는 코치이로 필터링
        if '코치 이름' in df.columns and '코치이 이름' in df.columns:
            filtered_df = df[
                (df['코치 이름'] == user_name) | 
                (df['코치이 이름'] == user_name)
            ]
        else:
            # 컬럼명이 다를 수 있으므로 첫 번째 텍스트 컬럼으로 시도
            filtered_df = df
    else:
        filtered_df = df
    
    if filtered_df.empty:
        st.info(f"{user_name}님과 관련된 1on1 코칭 기록이 없습니다.")
        return
    
    # 날짜 순으로 정렬 (최신순)
    date_col = None
    for col in ['코칭 날짜', '날짜', 'Date', 'date']:
        if col in filtered_df.columns:
            date_col = col
            break
    
    if date_col:
        try:
            filtered_df[date_col] = pd.to_datetime(filtered_df[date_col], errors='coerce')
            filtered_df = filtered_df.sort_values(by=date_col, ascending=False)
        except:
            pass
    
    # 필터링 옵션
    col1, col2 = st.columns(2)
    with col1:
        search_keyword = st.text_input(
            "🔍 검색",
            placeholder="주제, 참여자 이름 등으로 검색",
            help="코칭 주제나 참여자 이름으로 검색할 수 있습니다"
        )
    
    with col2:
        show_all = st.checkbox("전체 기록 보기", value=False)
    
    # 검색 필터 적용
    if search_keyword:
        keyword_lower = search_keyword.lower()
        mask = pd.Series([False] * len(filtered_df))
        for col in filtered_df.columns:
            mask = mask | filtered_df[col].astype(str).str.lower().str.contains(keyword_lower, na=False)
        filtered_df = filtered_df[mask]
    
    if not show_all and user_name:
        # 자신과 관련된 기록만 보기 (기본값)
        pass
    elif show_all:
        # 전체 기록 보기
        filtered_df = df
    
    # 기록 카드로 표시
    st.markdown("---")
    for idx, row in filtered_df.iterrows():
        with st.expander(
            f"📅 {row.get('코칭 날짜', row.get('날짜', '날짜 미지정'))} - {row.get('코칭 주제', row.get('주제', '주제 없음'))}",
            expanded=False
        ):
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**코치:** {row.get('코치 이름', 'N/A')}")
            with col2:
                st.markdown(f"**코치이:** {row.get('코치이 이름', 'N/A')}")
            
            st.markdown("---")
            
            if row.get('주요 내용') or row.get('내용'):
                st.markdown("### 주요 내용")
                st.write(row.get('주요 내용', row.get('내용', '')))
            
            if row.get('액션 아이템'):
                st.markdown("### 액션 아이템")
                st.write(row.get('액션 아이템', ''))
            
            if row.get('다음 미팅 일정'):
                st.markdown(f"**다음 미팅:** {row.get('다음 미팅 일정', '')}")
            
            if row.get('코칭 질 평가'):
                st.markdown(f"**코칭 질 평가:** {'⭐' * int(row.get('코칭 질 평가', 0))} ({row.get('코칭 질 평가', 0)}/5)")
            
            if row.get('기타 메모'):
                st.markdown("### 기타 메모")
                st.write(row.get('기타 메모', ''))

def format_cache_data_for_prompt(cache_data, data_type):
    """캐시 데이터를 프롬프트용 텍스트로 포맷팅합니다."""
    if not cache_data:
        return "데이터 없음"
    
    if isinstance(cache_data, list):
        if len(cache_data) == 0:
            return "데이터 없음"
        # 리스트의 각 레코드를 텍스트로 변환
        formatted = []
        for record in cache_data:
            if isinstance(record, dict):
                record_str = "\n".join([f"  - {k}: {v}" for k, v in record.items() if v and str(v).strip()])
                formatted.append(record_str)
        return "\n---\n".join(formatted)
    elif isinstance(cache_data, dict):
        return "\n".join([f"  - {k}: {v}" for k, v in cache_data.items() if v and str(v).strip()])
    else:
        return str(cache_data)

def create_word_document_from_feedback(feedback_text, title):
    """피드백 텍스트를 Word 문서로 변환합니다."""
    doc = Document()
    
    # 한글 폰트 설정 - 문서의 기본 스타일에 한글 폰트 적용
    try:
        # Normal 스타일 설정
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '맑은 고딕'
        normal_font.size = Pt(11)
        normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        # Heading 스타일들에도 한글 폰트 설정
        for i in range(1, 10):
            try:
                heading_style = doc.styles[f'Heading {i}']
                heading_font = heading_style.font
                heading_font.name = '맑은 고딕'
                heading_font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            except:
                pass
    except Exception as e:
        # 폰트 설정 실패 시 계속 진행
        pass
    
    # 제목 추가
    heading = doc.add_heading(title, 0)
    heading.alignment = 1  # 가운데 정렬
    
    # 제목 폰트 설정
    try:
        for run in heading.runs:
            run.font.name = '맑은 고딕'
            run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            run.font.size = Pt(20)
            run.font.bold = True
    except:
        pass
    
    # 날짜 추가
    date_para = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}")
    date_para.alignment = 1  # 가운데 정렬
    
    # 날짜 폰트 설정
    try:
        for run in date_para.runs:
            run.font.name = '맑은 고딕'
            run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    except:
        pass
    
    # 빈 줄 추가
    doc.add_paragraph()
    
    # 피드백 텍스트 파싱 및 추가
    lines = feedback_text.split('\n')
    current_para = None
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 빈 줄 처리
        if not line:
            doc.add_paragraph()
            current_para = None
            i += 1
            continue
        
        # 표 감지 (Markdown 표 형식: | 컬럼1 | 컬럼2 | 컬럼3 |)
        if line.startswith('|') and line.endswith('|'):
            # 표 시작 감지
            table_rows = []
            header_line = line
            separator_line = None
            
            # 헤더 다음 줄이 구분선인지 확인
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line.startswith('|') and ('---' in next_line or '===' in next_line or re.match(r'^\|\s*[-:=]+\s*\|', next_line)):
                    separator_line = next_line
                    i += 1  # 구분선 건너뛰기
            
            # 헤더 행 추가
            header_cells = [cell.strip() for cell in header_line.split('|')[1:-1]]
            table_rows.append(header_cells)
            
            # 데이터 행 수집
            i += 1
            while i < len(lines):
                current_line = lines[i].strip()
                if current_line.startswith('|') and current_line.endswith('|'):
                    cells = [cell.strip() for cell in current_line.split('|')[1:-1]]
                    if len(cells) == len(header_cells):  # 컬럼 수가 일치하는 경우에만
                        table_rows.append(cells)
                        i += 1
                    else:
                        break
                else:
                    break
            
            # 표를 Word 문서에 추가
            if len(table_rows) > 0:
                try:
                    num_cols = len(table_rows[0])
                    num_rows = len(table_rows)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Light Grid Accent 1'  # 표 스타일
                    
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            
                            # 셀 폰트 설정
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = '맑은 고딕'
                                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                            
                            # 헤더 행은 굵게
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                except Exception as e:
                    # 표 생성 실패 시 일반 텍스트로 처리
                    for row_data in table_rows:
                        doc.add_paragraph(' | '.join(row_data))
                        try:
                            para = doc.paragraphs[-1]
                            for run in para.runs:
                                run.font.name = '맑은 고딕'
                                run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                        except:
                            pass
            current_para = None
            i += 1
            continue
        
        # 일반 줄 처리
        # 제목 패턴 확인 (###, ##, #, 또는 숫자. 패턴)
        if re.match(r'^#{1,3}\s+.+', line):
            # Markdown 헤딩
            level = len(re.match(r'^(#{1,3})', line).group(1))
            text = re.sub(r'^#{1,3}\s+', '', line)
            heading_obj = None
            if level == 1:
                heading_obj = doc.add_heading(text, level=1)
            elif level == 2:
                heading_obj = doc.add_heading(text, level=2)
            else:
                heading_obj = doc.add_heading(text, level=3)
            # 헤딩에 한글 폰트 설정
            try:
                for run in heading_obj.runs:
                    run.font.name = '맑은 고딕'
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            except:
                pass
            current_para = None
            i += 1
        elif re.match(r'^\d+[\.\)]\s+.+', line):
            # 번호 목록
            para = doc.add_paragraph(line, style='List Number')
            # 번호 목록에 한글 폰트 설정
            try:
                for run in para.runs:
                    run.font.name = '맑은 고딕'
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            except:
                pass
            current_para = None
            i += 1
        elif re.match(r'^[-*]\s+.+', line):
            # 불릿 목록
            text = re.sub(r'^[-*]\s+', '', line)
            para = doc.add_paragraph(text, style='List Bullet')
            # 불릿 목록에 한글 폰트 설정
            try:
                for run in para.runs:
                    run.font.name = '맑은 고딕'
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            except:
                pass
            current_para = None
            i += 1
        elif re.match(r'^[•·]\s+.+', line):
            # 다른 불릿 문자
            text = re.sub(r'^[•·]\s+', '', line)
            para = doc.add_paragraph(text, style='List Bullet')
            # 불릿 목록에 한글 폰트 설정
            try:
                for run in para.runs:
                    run.font.name = '맑은 고딕'
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            except:
                pass
            current_para = None
        else:
            # 일반 텍스트
            if current_para is None:
                current_para = doc.add_paragraph(line)
                # 일반 텍스트에 한글 폰트 설정
                try:
                    for run in current_para.runs:
                        run.font.name = '맑은 고딕'
                        run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                except:
                    pass
            else:
                run = current_para.add_run(f"\n{line}")
                # 추가된 run에 한글 폰트 설정
                try:
                    run.font.name = '맑은 고딕'
                    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                except:
                    pass
            i += 1
    
    # 문서를 바이트로 변환
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def get_performance_coaching_feedback():
    """성과 코칭 피드백을 생성합니다."""
    try:
        if not GEMINI_API_KEY:
            return None, "Gemini API 키가 설정되지 않았습니다. .env 파일에 GEMINI_API_KEY를 설정해주세요."
        
        # 캐시 데이터 가져오기
        prefetch_cache = st.session_state.get('prefetch_cache') or {}
        if not isinstance(prefetch_cache, dict):
            prefetch_cache = {}
        
        archive_data = prefetch_cache.get('archive', [])
        mission_kpi_data = prefetch_cache.get('mission_kpi', [])
        ground_rule_data = prefetch_cache.get('ground_rule', [])
        
        # 데이터 확인
        if not archive_data and not mission_kpi_data and not ground_rule_data:
            return None, "피드백 생성을 위한 데이터가 없습니다. 먼저 데이터를 로드해주세요."
        
        # 프롬프트 구성
        prompt = f"""다음 정보를 바탕으로 업무 성과 창출과 관련한 상세한 코칭 피드백을 제공해주세요.

**사용자의 Snippet 아카이브:**
{format_cache_data_for_prompt(archive_data, 'archive')}

**조직의 Mission & KPI:**
{format_cache_data_for_prompt(mission_kpi_data, 'mission_kpi')}

**Team Ground Rule:**
{format_cache_data_for_prompt(ground_rule_data, 'ground_rule')}

위 정보를 종합적으로 분석하여 다음 항목들을 포함한 성과 코칭 피드백을 제공해주세요:
1. 현재 업무 성과의 강점 분석
2. 조직 목표와의 연계성 평가
3. 성과 개선을 위한 구체적인 액션 아이템
4. 팀 규칙과의 일치도 및 개선점
5. 향후 성과 창출을 위한 조언

피드백은 구체적이고 실행 가능한 내용으로 작성해주세요. 한국어로 답변해주세요."""

        # Gemini API 호출
        # 사용 가능한 모델 목록 확인
        model = None
        model_names_to_try = []
        last_error = None
        response = None
        
        # 먼저 사용 가능한 모델 목록 가져오기
        try:
            available_models = genai.list_models()
            # generateContent를 지원하고 Computer Use가 필요없는 모델 찾기
            for m in available_models:
                model_name = m.name
                # models/ 접두사 제거
                if model_name.startswith('models/'):
                    model_name = model_name.replace('models/', '')
                
                # generateContent를 지원하는 모델만
                if hasattr(m, 'supported_generation_methods') and 'generateContent' in m.supported_generation_methods:
                    # Computer Use 관련 모델 제외 (exp, 2.0-exp 등)
                    if 'exp' not in model_name.lower() and '2.0' not in model_name.lower():
                        if 'flash' in model_name.lower():
                            model_names_to_try.insert(0, model_name)  # flash 모델 우선
                        elif 'pro' in model_name.lower():
                            model_names_to_try.append(model_name)
                        else:
                            model_names_to_try.append(model_name)
        except Exception as e:
            # ListModels 실패 시 기본 모델 시도
            model_names_to_try = [
                'gemini-1.5-flash',
                'gemini-1.5-pro',
                'gemini-pro'
            ]
        
        # 만약 목록이 비어있으면 기본 모델 추가
        if not model_names_to_try:
            model_names_to_try = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
        
        # 모델명 시도 (Computer Use 오류 감지 및 건너뛰기)
        for model_name in model_names_to_try:
            try:
                model = genai.GenerativeModel(model_name)
                # Computer Use 없이 텍스트 생성만 시도
                response = model.generate_content(
                    prompt,
                    generation_config={
                        "temperature": 0.7,
                        "top_p": 0.8,
                        "top_k": 40,
                    }
                )
                # 성공적으로 응답을 받았으면 이 모델 사용
                if response:
                    break
            except Exception as e:
                error_msg = str(e)
                last_error = error_msg
                # Computer Use 관련 오류는 이 모델을 건너뛰기
                if 'Computer Use' in error_msg or 'computer-use' in error_msg.lower():
                    model = None
                    response = None
                    continue
                # 404 오류는 다른 모델 시도
                if '404' in error_msg or 'not found' in error_msg.lower():
                    model = None
                    response = None
                    continue
                # 다른 오류도 일단 건너뛰고 다음 모델 시도
                model = None
                response = None
                continue
        
        if model is None or response is None:
            # 사용 가능한 모델 목록 정보 추가
            available_info = ""
            try:
                available_models = genai.list_models()
                available_names = [m.name for m in available_models if hasattr(m, 'supported_generation_methods') and 'generateContent' in m.supported_generation_methods]
                if available_names:
                    available_info = f" 사용 가능한 모델: {available_names[:5]}"
            except:
                pass
            raise Exception(f"사용 가능한 Gemini 모델을 찾을 수 없습니다. 마지막 오류: {last_error}. 시도한 모델: {model_names_to_try}.{available_info} API 키와 모델 접근 권한을 확인해주세요.")
        
        # 응답 텍스트 추출
        if hasattr(response, 'text'):
            feedback_text = response.text
        elif hasattr(response, 'candidates') and len(response.candidates) > 0:
            if hasattr(response.candidates[0], 'content'):
                if hasattr(response.candidates[0].content, 'parts'):
                    feedback_text = response.candidates[0].content.parts[0].text
                else:
                    feedback_text = str(response.candidates[0].content)
            else:
                feedback_text = str(response.candidates[0])
        else:
            feedback_text = str(response)
        
        return feedback_text, None
        
    except Exception as e:
        return None, f"피드백 생성 중 오류가 발생했습니다: {str(e)}"

def get_growth_coaching_feedback():
    """성장 코칭 피드백을 생성합니다."""
    try:
        if not GEMINI_API_KEY:
            return None, "Gemini API 키가 설정되지 않았습니다. .env 파일에 GEMINI_API_KEY를 설정해주세요."
        
        # 캐시 데이터 가져오기
        prefetch_cache = st.session_state.get('prefetch_cache', {})
        
        archive_data = prefetch_cache.get('archive', [])
        cdp_data = prefetch_cache.get('cdp', [])
        idp_data = prefetch_cache.get('idp', [])
        mission_kpi_data = prefetch_cache.get('mission_kpi', [])
        ground_rule_data = prefetch_cache.get('ground_rule', [])
        
        # 데이터 확인
        if not archive_data and not cdp_data and not idp_data and not mission_kpi_data and not ground_rule_data:
            return None, "피드백 생성을 위한 데이터가 없습니다. 먼저 데이터를 로드해주세요."
        
        # 프롬프트 구성
        prompt = f"""다음 정보를 바탕으로 성장과 관련한 상세한 코칭 피드백을 제공해주세요.

**사용자의 Snippet 아카이브:**
{format_cache_data_for_prompt(archive_data, 'archive')}

**경력 개발 계획 (CDP):**
{format_cache_data_for_prompt(cdp_data, 'cdp')}

**개인 개발 계획 (IDP):**
{format_cache_data_for_prompt(idp_data, 'idp')}

**조직의 Mission & KPI:**
{format_cache_data_for_prompt(mission_kpi_data, 'mission_kpi')}

**Team Ground Rule:**
{format_cache_data_for_prompt(ground_rule_data, 'ground_rule')}

위 정보를 종합적으로 분석하여 다음 항목들을 포함한 성장 코칭 피드백을 제공해주세요:
1. 현재 성장 상태와 역량 분석
2. 개인 개발 계획(IDP) 및 경력 개발 계획(CDP) 달성도 평가
3. 조직 목표와의 정렬도 및 성장 방향성 제시
4. 성장을 위한 구체적인 학습 및 개발 액션 아이템
5. 다음 단계 성장을 위한 조언 및 로드맵

피드백은 구체적이고 실행 가능한 내용으로 작성해주세요. 한국어로 답변해주세요."""

        # Gemini API 호출
        # 사용 가능한 모델 목록 확인
        model = None
        model_names_to_try = []
        last_error = None
        response = None
        
        # 먼저 사용 가능한 모델 목록 가져오기
        try:
            available_models = genai.list_models()
            # generateContent를 지원하고 Computer Use가 필요없는 모델 찾기
            for m in available_models:
                model_name = m.name
                # models/ 접두사 제거
                if model_name.startswith('models/'):
                    model_name = model_name.replace('models/', '')
                
                # generateContent를 지원하는 모델만
                if hasattr(m, 'supported_generation_methods') and 'generateContent' in m.supported_generation_methods:
                    # Computer Use 관련 모델 제외 (exp, 2.0-exp 등)
                    if 'exp' not in model_name.lower() and '2.0' not in model_name.lower():
                        if 'flash' in model_name.lower():
                            model_names_to_try.insert(0, model_name)  # flash 모델 우선
                        elif 'pro' in model_name.lower():
                            model_names_to_try.append(model_name)
                        else:
                            model_names_to_try.append(model_name)
        except Exception as e:
            # ListModels 실패 시 기본 모델 시도
            model_names_to_try = [
                'gemini-1.5-flash',
                'gemini-1.5-pro',
                'gemini-pro'
            ]
        
        # 만약 목록이 비어있으면 기본 모델 추가
        if not model_names_to_try:
            model_names_to_try = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
        
        # 모델명 시도 (Computer Use 오류 감지 및 건너뛰기)
        for model_name in model_names_to_try:
            try:
                model = genai.GenerativeModel(model_name)
                # Computer Use 없이 텍스트 생성만 시도
                response = model.generate_content(
                    prompt,
                    generation_config={
                        "temperature": 0.7,
                        "top_p": 0.8,
                        "top_k": 40,
                    }
                )
                # 성공적으로 응답을 받았으면 이 모델 사용
                if response:
                    break
            except Exception as e:
                error_msg = str(e)
                last_error = error_msg
                # Computer Use 관련 오류는 이 모델을 건너뛰기
                if 'Computer Use' in error_msg or 'computer-use' in error_msg.lower():
                    model = None
                    response = None
                    continue
                # 404 오류는 다른 모델 시도
                if '404' in error_msg or 'not found' in error_msg.lower():
                    model = None
                    response = None
                    continue
                # 다른 오류도 일단 건너뛰고 다음 모델 시도
                model = None
                response = None
                continue
        
        if model is None or response is None:
            # 사용 가능한 모델 목록 정보 추가
            available_info = ""
            try:
                available_models = genai.list_models()
                available_names = [m.name for m in available_models if hasattr(m, 'supported_generation_methods') and 'generateContent' in m.supported_generation_methods]
                if available_names:
                    available_info = f" 사용 가능한 모델: {available_names[:5]}"
            except:
                pass
            raise Exception(f"사용 가능한 Gemini 모델을 찾을 수 없습니다. 마지막 오류: {last_error}. 시도한 모델: {model_names_to_try}.{available_info} API 키와 모델 접근 권한을 확인해주세요.")
        
        # 응답 텍스트 추출
        if hasattr(response, 'text'):
            feedback_text = response.text
        elif hasattr(response, 'candidates') and len(response.candidates) > 0:
            if hasattr(response.candidates[0], 'content'):
                if hasattr(response.candidates[0].content, 'parts'):
                    feedback_text = response.candidates[0].content.parts[0].text
                else:
                    feedback_text = str(response.candidates[0].content)
            else:
                feedback_text = str(response.candidates[0])
        else:
            feedback_text = str(response)
        
        return feedback_text, None
        
    except Exception as e:
        return None, f"피드백 생성 중 오류가 발생했습니다: {str(e)}"

def render_performance_feedback():
    """성과 코칭 피드백을 렌더링합니다."""
    st.subheader("📊 성과 코칭 피드백")
    st.markdown("사용자의 업무 성과를 분석하여 코칭 피드백을 제공합니다.")
    
    # 데이터 상태 표시
    prefetch_cache = st.session_state.get('prefetch_cache') or {}
    if not isinstance(prefetch_cache, dict):
        prefetch_cache = {}
    archive_data = prefetch_cache.get('archive', [])
    mission_kpi_data = prefetch_cache.get('mission_kpi', [])
    ground_rule_data = prefetch_cache.get('ground_rule', [])
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Snippet 아카이브", f"{len(archive_data)}개" if archive_data else "없음")
    with col2:
        st.metric("Mission & KPI", f"{len(mission_kpi_data)}개" if mission_kpi_data else "없음")
    with col3:
        st.metric("Team Ground Rule", f"{len(ground_rule_data)}개" if ground_rule_data else "없음")
    
    if not GEMINI_API_KEY:
        st.warning("⚠️ Gemini API 키가 설정되지 않았습니다. .env 파일에 GEMINI_API_KEY를 설정해주세요.")
    
    if st.button("🔍 성과 피드백 생성하기", use_container_width=True, type="primary"):
        with st.spinner("AI가 피드백을 생성하는 중입니다. 잠시만 기다려주세요..."):
            feedback, error = get_performance_coaching_feedback()
            
            if error:
                st.error(error)
            elif feedback:
                st.success("✅ 성과 코칭 피드백이 생성되었습니다!")
                st.markdown("---")
                st.markdown("### 📋 코칭 피드백")
                st.markdown(feedback)
                
                # 피드백 다운로드 버튼 (Word 파일)
                word_bytes = create_word_document_from_feedback(feedback, "성과 코칭 피드백")
                st.download_button(
                    label="📥 피드백 다운로드 (Word)",
                    data=word_bytes,
                    file_name=f"성과_코칭_피드백_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("피드백을 생성할 수 없습니다.")

def render_performance_feedback_auto():
    """성과 코칭 피드백을 자동으로 생성하고 표시합니다."""
    with st.expander("📊 성과 코칭 피드백", expanded=True):
        st.markdown("사용자의 업무 성과를 분석하여 코칭 피드백을 제공합니다.")
        
        feedback_placeholder = st.empty()
        with feedback_placeholder:
            with st.spinner("🤖 성과 코칭 피드백 생성 중..."):
                feedback, error = get_performance_coaching_feedback()
        
        feedback_placeholder.empty()
        
        if error:
            st.error(f"❌ {error}")
        elif feedback:
            st.markdown("### 📋 코칭 피드백")
            st.markdown(feedback)
            
            # 피드백 다운로드 버튼 (Word 파일)
            word_bytes = create_word_document_from_feedback(feedback, "성과 코칭 피드백")
            st.download_button(
                label="📥 피드백 다운로드 (Word)",
                data=word_bytes,
                file_name=f"성과_코칭_피드백_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("피드백을 생성할 수 없습니다.")

def render_growth_feedback():
    """성장 코칭 피드백을 렌더링합니다."""
    st.subheader("🌱 성장 코칭 피드백")
    st.markdown("사용자의 성장 상황을 분석하여 코칭 피드백을 제공합니다.")
    
    # 데이터 상태 표시
    prefetch_cache = st.session_state.get('prefetch_cache') or {}
    if not isinstance(prefetch_cache, dict):
        prefetch_cache = {}
    archive_data = prefetch_cache.get('archive', [])
    cdp_data = prefetch_cache.get('cdp', [])
    idp_data = prefetch_cache.get('idp', [])
    mission_kpi_data = prefetch_cache.get('mission_kpi', [])
    ground_rule_data = prefetch_cache.get('ground_rule', [])
    
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.metric("Snippet", f"{len(archive_data)}개" if archive_data else "없음")
    with col2:
        st.metric("CDP", f"{len(cdp_data)}개" if cdp_data else "없음")
    with col3:
        st.metric("IDP", f"{len(idp_data)}개" if idp_data else "없음")
    with col4:
        st.metric("Mission & KPI", f"{len(mission_kpi_data)}개" if mission_kpi_data else "없음")
    with col5:
        st.metric("Ground Rule", f"{len(ground_rule_data)}개" if ground_rule_data else "없음")
    
    if not GEMINI_API_KEY:
        st.warning("⚠️ Gemini API 키가 설정되지 않았습니다. .env 파일에 GEMINI_API_KEY를 설정해주세요.")
    
    if st.button("🔍 성장 피드백 생성하기", use_container_width=True, type="primary"):
        with st.spinner("AI가 피드백을 생성하는 중입니다. 잠시만 기다려주세요..."):
            feedback, error = get_growth_coaching_feedback()
            
            if error:
                st.error(error)
            elif feedback:
                st.success("✅ 성장 코칭 피드백이 생성되었습니다!")
                st.markdown("---")
                st.markdown("### 📋 코칭 피드백")
                st.markdown(feedback)
                
                # 피드백 다운로드 버튼 (Word 파일)
                word_bytes = create_word_document_from_feedback(feedback, "성장 코칭 피드백")
                st.download_button(
                    label="📥 피드백 다운로드 (Word)",
                    data=word_bytes,
                    file_name=f"성장_코칭_피드백_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("피드백을 생성할 수 없습니다.")

def render_growth_feedback_auto():
    """성장 코칭 피드백을 자동으로 생성하고 표시합니다."""
    with st.expander("🌱 성장 코칭 피드백", expanded=True):
        st.markdown("사용자의 성장 상황을 분석하여 코칭 피드백을 제공합니다.")
        
        feedback_placeholder = st.empty()
        with feedback_placeholder:
            with st.spinner("🤖 성장 코칭 피드백 생성 중..."):
                feedback, error = get_growth_coaching_feedback()
        
        feedback_placeholder.empty()
        
        if error:
            st.error(f"❌ {error}")
        elif feedback:
            st.markdown("### 📋 코칭 피드백")
            st.markdown(feedback)
            
            # 피드백 다운로드 버튼 (Word 파일)
            word_bytes = create_word_document_from_feedback(feedback, "성장 코칭 피드백")
            st.download_button(
                label="📥 피드백 다운로드 (Word)",
                data=word_bytes,
                file_name=f"성장_코칭_피드백_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("피드백을 생성할 수 없습니다.")

def ensure_cache_data():
    """필요한 캐시 데이터가 있는지 확인하고, 없으면 로드합니다."""
    # prefetch_cache 초기화
    if 'prefetch_cache' not in st.session_state:
        st.session_state.prefetch_cache = {}
    
    prefetch_cache = st.session_state.prefetch_cache
    if not isinstance(prefetch_cache, dict):
        prefetch_cache = {}
        st.session_state.prefetch_cache = prefetch_cache
    
    user_info = st.session_state.get('user_info')
    if not user_info:
        return False
    
    user_name = user_info.get('name')
    if not user_name:
        return False
    
    missing_data = []
    need_load = False
    
    # 캐시 확인
    if not prefetch_cache.get('archive'):
        missing_data.append('Snippet 아카이브')
        need_load = True
    if not prefetch_cache.get('cdp'):
        missing_data.append('CDP')
        need_load = True
    if not prefetch_cache.get('idp'):
        missing_data.append('IDP')
        need_load = True
    if not prefetch_cache.get('mission_kpi'):
        missing_data.append('Mission & KPI')
        need_load = True
    if not prefetch_cache.get('ground_rule'):
        missing_data.append('Team Ground Rule')
        need_load = True
    
    if not need_load:
        return True
    
    # 데이터 로딩
    status_text = st.empty()
    progress_bar = st.progress(0)
    
    # 1. Snippet 아카이브 로딩
    if not prefetch_cache.get('archive'):
        status_text.info("📚 Snippet 아카이브 데이터 로딩 중...")
        progress_bar.progress(10)
        try:
            import sys
            import Archive
            
            # main.py 모듈 가져오기 (여러 방법 시도)
            main_module = None
            get_client = None
            spreadsheet_id = None
            
            # 방법 1: sys.modules에서 찾기
            main_module = sys.modules.get('main')
            
            # 방법 2: 직접 import 시도
            if not main_module:
                try:
                    import main as main_mod
                    main_module = main_mod
                except Exception:
                    pass
            
            # 방법 3: importlib로 로드 시도
            if not main_module:
                try:
                    import importlib
                    main_module = importlib.import_module('main')
                except Exception:
                    pass
            
            # main.py의 함수들을 직접 사용하여 로드
            main_get_client = None
            main_spreadsheet_id = None
            
            if main_module:
                main_get_client = getattr(main_module, 'get_google_sheets_client', None)
                main_spreadsheet_id = getattr(main_module, 'SPREADSHEET_ID', None)
            
            # 데이터 로드 시도
            archive_df = None
            
            # 디버깅 정보 수집
            debug_info = []
            debug_info.append(f"사용자 이름: {user_name}")
            debug_info.append(f"Google Sheets 연결: {st.session_state.get('google_sheets_connected', False)}")
            debug_info.append(f"main_module 찾기: {main_module is not None}")
            debug_info.append(f"main.py에서 get_client 찾기: {main_get_client is not None}")
            debug_info.append(f"main.py에서 SPREADSHEET_ID 찾기: {main_spreadsheet_id}")
            
            # 최종 값 결정 (우선순위: main.py > 1on1.py > 하드코딩)
            if main_get_client:
                get_client = main_get_client
                debug_info.append("✅ main.py의 get_google_sheets_client 사용")
            else:
                get_client = get_google_sheets_client
                debug_info.append("✅ 1on1.py의 get_google_sheets_client 사용")
            
            if main_spreadsheet_id:
                spreadsheet_id = main_spreadsheet_id
                debug_info.append(f"✅ main.py의 SPREADSHEET_ID 사용: {spreadsheet_id}")
            else:
                spreadsheet_id = "1THmwStR6p0_SUyLEV6-edT0kigANvTCPOkAzN7NaEQE"
                debug_info.append(f"✅ 하드코딩된 spreadsheet_id 사용: {spreadsheet_id}")
            
            debug_info.append(f"최종 get_client: {get_client is not None} (타입: {type(get_client).__name__})")
            debug_info.append(f"최종 spreadsheet_id: {spreadsheet_id}")
            
            # 1순위: Google Sheets에서 로드
            if get_client and spreadsheet_id:
                # google_sheets_connected 상태와 무관하게 시도 (연결 상태가 잘못 표시될 수 있음)
                try:
                    debug_info.append(f"Google Sheets에서 로드 시도 중... (spreadsheet_id: {spreadsheet_id})")
                    # get_client가 함수인지 확인
                    if callable(get_client):
                        debug_info.append("get_client는 호출 가능한 함수입니다")
                        archive_df = Archive.get_snippets_from_google_sheets(get_client, spreadsheet_id)
                    else:
                        debug_info.append(f"get_client가 함수가 아닙니다: {type(get_client)}")
                        # 직접 시도
                        client = None
                        if main_module and hasattr(main_module, 'get_google_sheets_client'):
                            client = main_module.get_google_sheets_client()
                            if client:
                                spreadsheet = client.open_by_key(spreadsheet_id)
                                worksheet = spreadsheet.worksheet("Sheet1")
                                records = worksheet.get_all_records()
                                archive_df = pd.DataFrame(records)
                            else:
                                archive_df = None
                        else:
                            archive_df = None
                    
                    if archive_df is not None and not archive_df.empty:
                        debug_info.append(f"Google Sheets 로드 성공: {len(archive_df)}개 행, {len(archive_df.columns)}개 컬럼")
                    else:
                        debug_info.append("Google Sheets 데이터가 비어있거나 None")
                except Exception as gs_error:
                    import traceback
                    error_trace = traceback.format_exc()
                    debug_info.append(f"Google Sheets 로딩 실패: {str(gs_error)}")
                    debug_info.append(f"상세 오류: {error_trace[:500]}")  # 처음 500자만
                    archive_df = None
            else:
                debug_info.append(f"조건 불만족 - get_client: {get_client is not None}, spreadsheet_id: {spreadsheet_id}")
                if not get_client:
                    debug_info.append("get_client 함수를 찾을 수 없습니다")
                if not spreadsheet_id:
                    debug_info.append("spreadsheet_id를 찾을 수 없습니다")
            
            # 2순위: 로컬 CSV에서 로드
            if archive_df is None or (hasattr(archive_df, 'empty') and archive_df.empty):
                try:
                    debug_info.append("로컬 CSV에서 로드 시도 중...")
                    archive_df = Archive.get_snippets_from_local_csv()
                    if archive_df is not None and not archive_df.empty:
                        debug_info.append(f"로컬 CSV 로드 성공: {len(archive_df)}개 행")
                    else:
                        debug_info.append("로컬 CSV 데이터가 비어있거나 파일 없음")
                except Exception as csv_error:
                    debug_info.append(f"로컬 CSV 로딩 실패: {str(csv_error)}")
                    archive_df = None
            
            # 데이터 처리
            if archive_df is not None and not archive_df.empty:
                debug_info.append(f"데이터프레임 컬럼: {list(archive_df.columns)}")
                
                # 컬럼명 확인 (대소문자 무시, 공백 제거)
                name_column = None
                for col in archive_df.columns:
                    col_clean = str(col).strip().lower()
                    if '이름' in str(col) or 'name' in col_clean:
                        name_column = col
                        debug_info.append(f"이름 컬럼 찾음: {col}")
                        break
                
                if name_column:
                    # 사용자 이름으로 필터링
                    user_archive = archive_df[archive_df[name_column] == user_name]
                    debug_info.append(f"사용자 '{user_name}' 매칭 결과: {len(user_archive)}개 행")
                    
                    # 정확히 매칭되지 않으면 부분 매칭 시도
                    if user_archive.empty:
                        # 부분 매칭 시도 (공백 제거)
                        user_name_clean = str(user_name).strip()
                        for idx, row in archive_df.iterrows():
                            row_name = str(row[name_column]).strip() if pd.notna(row[name_column]) else ""
                            if user_name_clean == row_name:
                                user_archive = archive_df[archive_df.index == idx]
                                debug_info.append(f"부분 매칭 성공: 인덱스 {idx}")
                                break
                    
                    prefetch_cache['archive'] = user_archive.to_dict('records') if not user_archive.empty else []
                else:
                    debug_info.append("이름 컬럼을 찾을 수 없음 - 전체 데이터 사용")
                    # 이름 컬럼이 없으면 전체 데이터 사용
                    prefetch_cache['archive'] = archive_df.to_dict('records')
            else:
                debug_info.append("데이터프레임이 비어있거나 None")
                prefetch_cache['archive'] = []
           
            
            # 로딩 실패 시 경고 표시
            if not prefetch_cache.get('archive'):
                if debug_info:
                    # 마지막 몇 개 디버그 메시지만 표시
                    recent_errors = [d for d in debug_info if any(keyword in d for keyword in ['실패', '없음', '비어', 'None', '불만족', '찾을 수 없'])]
                    if recent_errors:
                        error_msg = " | ".join(recent_errors[-3:])
                        st.error(f"❌ Snippet 아카이브 데이터를 찾을 수 없습니다.\n\n**오류 정보:**\n{error_msg}\n\n**해결 방법:**\n- 위의 '🔍 Snippet 아카이브 로딩 상세 정보'를 펼쳐서 자세한 정보를 확인하세요.\n- Google Sheets 연결 상태를 확인하세요.\n- 사용자 이름이 데이터의 '이름' 컬럼과 정확히 일치하는지 확인하세요.")
                
        except Exception as e:
            # 에러 상세 정보를 로그로 남기기
            import traceback
            error_detail = traceback.format_exc()
            st.warning(f"⚠️ Snippet 아카이브 로딩 중 오류: {str(e)}")
            # 디버그 정보는 개발 환경에서만 표시
            if st.session_state.get('debug_mode', False):
                st.text(f"상세 오류:\n{error_detail}")
            prefetch_cache['archive'] = []
    
    # 2. CDP 로딩
    if not prefetch_cache.get('cdp'):
        status_text.info("📊 CDP 데이터 로딩 중...")
        progress_bar.progress(30)
        try:
            import cdp
            cdp_df = cdp._fetch_cdp_dataframe()
            if cdp_df is not None and not cdp_df.empty:
                normalized = {c.strip(): c for c in cdp_df.columns}
                name_col = normalized.get("이름") or normalized.get("name") or list(cdp_df.columns)[0]
                user_cdp = cdp_df[cdp_df[name_col] == user_name]
                prefetch_cache['cdp'] = user_cdp.to_dict('records') if not user_cdp.empty else []
            else:
                prefetch_cache['cdp'] = []
        except Exception:
            prefetch_cache['cdp'] = []
    
    # 3. IDP 로딩
    if not prefetch_cache.get('idp'):
        status_text.info("🎯 IDP 데이터 로딩 중...")
        progress_bar.progress(50)
        try:
            import idp_usage
            idp_df = idp_usage.fetch_idp_dataframe()
            if idp_df is not None and not idp_df.empty:
                if '이름' in idp_df.columns:
                    user_idp = idp_df[idp_df['이름'] == user_name]
                    prefetch_cache['idp'] = user_idp.to_dict('records') if not user_idp.empty else []
                else:
                    prefetch_cache['idp'] = idp_df.to_dict('records')
            else:
                prefetch_cache['idp'] = []
        except Exception:
            prefetch_cache['idp'] = []
    
    # 4. Mission & KPI 로딩
    if not prefetch_cache.get('mission_kpi'):
        status_text.info("🎯 Mission & KPI 데이터 로딩 중...")
        progress_bar.progress(70)
        try:
            import organization
            mission_kpi_df = organization.get_sheet_data(organization.MISSION_KPI_SHEET_ID)
            if mission_kpi_df is not None and not mission_kpi_df.empty:
                prefetch_cache['mission_kpi'] = mission_kpi_df.to_dict('records')
            else:
                prefetch_cache['mission_kpi'] = []
        except Exception:
            prefetch_cache['mission_kpi'] = []
    
    # 5. Team Ground Rule 로딩
    if not prefetch_cache.get('ground_rule'):
        status_text.info("📋 Team Ground Rule 데이터 로딩 중...")
        progress_bar.progress(90)
        try:
            import organization
            ground_rule_df = organization.get_sheet_data(organization.GROUND_RULE_SHEET_ID)
            if ground_rule_df is not None and not ground_rule_df.empty:
                prefetch_cache['ground_rule'] = ground_rule_df.to_dict('records')
            else:
                prefetch_cache['ground_rule'] = []
        except Exception:
            prefetch_cache['ground_rule'] = []
    
    # 캐시 저장
    st.session_state.prefetch_cache = prefetch_cache
    
    status_text.empty()
    progress_bar.empty()
    
    return True

def render_oneon1_embedded():
    """1on1 코칭 페이지를 임베드 모드로 렌더링합니다 (main.py에서 사용)."""
    st.title("👥 1on1 코칭")
    st.markdown("AI 기반 코칭 피드백을 제공합니다.")
    st.markdown("---")
    
    # 캐시 데이터 확인 및 로딩
    status_container = st.container()
    with status_container:
        status_placeholder = st.empty()
        progress_placeholder = st.empty()
        
        status_placeholder.info("📊 데이터 확인 중...")
        progress_placeholder.progress(0)
        
        # 데이터 로딩
        status_placeholder.info("📥 데이터 로딩 중...")
        progress_placeholder.progress(0.2)
        
        ensure_cache_data()
        
        # 데이터 분석 중
        status_placeholder.info("🔍 데이터 분석 중...")
        progress_placeholder.progress(0.5)
        
        # 캐시 상태 확인
        prefetch_cache = st.session_state.get('prefetch_cache') or {}
        if not isinstance(prefetch_cache, dict):
            prefetch_cache = {}
        
        archive_count = len(prefetch_cache.get('archive', []))
        cdp_count = len(prefetch_cache.get('cdp', []))
        idp_count = len(prefetch_cache.get('idp', []))
        mission_kpi_count = len(prefetch_cache.get('mission_kpi', []))
        ground_rule_count = len(prefetch_cache.get('ground_rule', []))
        
        # 데이터 요약 표시
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.metric("Snippet", f"{archive_count}개")
        with col2:
            st.metric("CDP", f"{cdp_count}개")
        with col3:
            st.metric("IDP", f"{idp_count}개")
        with col4:
            st.metric("Mission & KPI", f"{mission_kpi_count}개")
        with col5:
            st.metric("Ground Rule", f"{ground_rule_count}개")
        
        progress_placeholder.progress(0.7)
        status_placeholder.info("🤖 코칭 피드백 생성 중...")
        progress_placeholder.progress(0.8)
    
    # 상태 표시 제거
    status_placeholder.empty()
    progress_placeholder.empty()
    
    st.markdown("---")
    
    # 성과 코칭 피드백 (자동 생성) - 접이식 카드
    render_performance_feedback_auto()
    
    st.markdown("<br>", unsafe_allow_html=True)  # 카드 간 간격
    
    # 성장 코칭 피드백 (자동 생성) - 접이식 카드
    render_growth_feedback_auto()

def main():
    """독립 실행용 메인 함수"""
    st.set_page_config(
        page_title="1on1 코칭",
        page_icon="👥",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    render_oneon1_embedded()

if __name__ == "__main__":
    main()

